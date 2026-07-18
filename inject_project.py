import sys
import argparse
from docx import Document
from docx.shared import Pt, RGBColor

def clone_paragraph_with_text(src_p, dest_p, text):
    # Copy paragraph style
    dest_p.style = src_p.style
    
    # Copy paragraph format properties (spacing and alignment)
    dest_p.paragraph_format.space_before = src_p.paragraph_format.space_before
    dest_p.paragraph_format.space_after = src_p.paragraph_format.space_after
    dest_p.paragraph_format.line_spacing = src_p.paragraph_format.line_spacing
    dest_p.paragraph_format.line_spacing_rule = src_p.paragraph_format.line_spacing_rule
    dest_p.paragraph_format.alignment = src_p.paragraph_format.alignment
    
    # Copy paragraph indents (CRITICAL for column borders and margins)
    dest_p.paragraph_format.left_indent = src_p.paragraph_format.left_indent
    dest_p.paragraph_format.right_indent = src_p.paragraph_format.right_indent
    dest_p.paragraph_format.first_line_indent = src_p.paragraph_format.first_line_indent
    
    # Copy widow control and page behavior
    dest_p.paragraph_format.keep_with_next = src_p.paragraph_format.keep_with_next
    dest_p.paragraph_format.page_break_before = src_p.paragraph_format.page_break_before
    dest_p.paragraph_format.widow_control = src_p.paragraph_format.widow_control
    
    # Copy run formatting from the first run of source paragraph
    if src_p.runs:
        src_run = src_p.runs[0]
        dest_run = dest_p.add_run(text)
        dest_run.font.name = src_run.font.name
        dest_run.font.size = src_run.font.size
        dest_run.font.bold = src_run.font.bold
        dest_run.font.italic = src_run.font.italic
        if src_run.font.color and src_run.font.color.rgb:
            dest_run.font.color.rgb = src_run.font.color.rgb
    else:
        dest_p.text = text

def inject_project(doc_path, output_path, title, tech, description, column='auto', position='last'):
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error opening document {doc_path}: {e}", file=sys.stderr)
        return False

    if not doc.tables:
        print("Error: Document has no tables.", file=sys.stderr)
        return False
        
    table = doc.tables[0]
    
    # Select the target cell based on the column parameter
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    
    if column == 'left':
        target_cell = left_cell
        cell_name = 'left'
    elif column == 'right':
        target_cell = right_cell
        cell_name = 'right'
    else: # 'auto' - Choose the column with fewer paragraphs to balance layout
        left_len = len(left_cell.paragraphs)
        right_len = len(right_cell.paragraphs)
        print(f"Auto-balancing columns: Left column has {left_len} paragraphs, Right column has {right_len} paragraphs.")
        if left_len <= right_len:
            target_cell = left_cell
            cell_name = 'left'
        else:
            target_cell = right_cell
            cell_name = 'right'
            
    print(f"Selected target column: '{cell_name}'")

    # Determine insertion behavior
    if cell_name == 'left':
        # Left cell has multiple sections (MORE PROJECTS, DEVELOPER TOOLS & PACKAGES)
        if position == 'first':
            # Insert at the top of the projects list
            more_projects_idx = -1
            for idx, p in enumerate(target_cell.paragraphs):
                if "MORE PROJECTS" in p.text:
                    more_projects_idx = idx
                    break
            
            if more_projects_idx == -1:
                print("Error: 'MORE PROJECTS' header not found in left cell.", file=sys.stderr)
                return False
                
            title_idx = -1
            for idx in range(more_projects_idx + 1, len(target_cell.paragraphs)):
                if target_cell.paragraphs[idx].text.strip():
                    title_idx = idx
                    break
                    
            if title_idx == -1 or title_idx + 2 >= len(target_cell.paragraphs):
                print("Error: Could not locate existing project templates in left cell.", file=sys.stderr)
                return False
                
            p_title_src = target_cell.paragraphs[title_idx]
            p_tech_src = target_cell.paragraphs[title_idx + 1]
            p_desc_src = target_cell.paragraphs[title_idx + 2]
            
            # Insert before the first project
            p_desc_new = p_title_src.insert_paragraph_before()
            clone_paragraph_with_text(p_desc_src, p_desc_new, description)
            p_tech_new = p_desc_new.insert_paragraph_before()
            clone_paragraph_with_text(p_tech_src, p_tech_new, tech)
            p_title_new = p_tech_new.insert_paragraph_before()
            clone_paragraph_with_text(p_title_src, p_title_new, title)
        else:
            # position == 'last' - Insert at the end of the projects list (before DEVELOPER TOOLS)
            next_header_idx = -1
            for idx, p in enumerate(target_cell.paragraphs):
                if "DEVELOPER TOOLS" in p.text:
                    next_header_idx = idx
                    break
            
            if next_header_idx == -1:
                # If no DEVELOPER TOOLS section, append to end of cell
                p_title_src = target_cell.paragraphs[-3]
                p_tech_src = target_cell.paragraphs[-2]
                p_desc_src = target_cell.paragraphs[-1]
                
                p_title_new = target_cell.add_paragraph()
                clone_paragraph_with_text(p_title_src, p_title_new, title)
                p_tech_new = target_cell.add_paragraph()
                clone_paragraph_with_text(p_tech_src, p_tech_new, tech)
                p_desc_new = target_cell.add_paragraph()
                clone_paragraph_with_text(p_desc_src, p_desc_new, description)
            else:
                # Check for empty line before next header
                if next_header_idx > 0 and not target_cell.paragraphs[next_header_idx - 1].text.strip():
                    target_p_idx = next_header_idx - 1
                else:
                    target_p_idx = next_header_idx
                    
                target_p = target_cell.paragraphs[target_p_idx]
                
                p_title_src = target_cell.paragraphs[target_p_idx - 3]
                p_tech_src = target_cell.paragraphs[target_p_idx - 2]
                p_desc_src = target_cell.paragraphs[target_p_idx - 1]
                
                p_desc_new = target_p.insert_paragraph_before()
                clone_paragraph_with_text(p_desc_src, p_desc_new, description)
                p_tech_new = p_desc_new.insert_paragraph_before()
                clone_paragraph_with_text(p_tech_src, p_tech_new, tech)
                p_title_new = p_tech_new.insert_paragraph_before()
                clone_paragraph_with_text(p_title_src, p_title_new, title)
    else:
        # Right cell (only contains projects)
        if position == 'first':
            # Insert at the top of the projects list (index 2, after "MORE PROJECTS" header)
            p_title_src = target_cell.paragraphs[2]
            p_tech_src = target_cell.paragraphs[3]
            p_desc_src = target_cell.paragraphs[4]
            
            p_desc_new = p_title_src.insert_paragraph_before()
            clone_paragraph_with_text(p_desc_src, p_desc_new, description)
            p_tech_new = p_desc_new.insert_paragraph_before()
            clone_paragraph_with_text(p_tech_src, p_tech_new, tech)
            p_title_new = p_tech_new.insert_paragraph_before()
            clone_paragraph_with_text(p_title_src, p_title_new, title)
        else:
            # position == 'last' - Append to the very end of the right column
            p_title_src = target_cell.paragraphs[-3]
            p_tech_src = target_cell.paragraphs[-2]
            p_desc_src = target_cell.paragraphs[-1]
            
            p_title_new = target_cell.add_paragraph()
            clone_paragraph_with_text(p_title_src, p_title_new, title)
            p_tech_new = target_cell.add_paragraph()
            clone_paragraph_with_text(p_tech_src, p_tech_new, tech)
            p_desc_new = target_cell.add_paragraph()
            clone_paragraph_with_text(p_desc_src, p_desc_new, description)

    try:
        doc.save(output_path)
        print(f"Successfully injected project into {output_path} (Column: {cell_name}, Position: {position})")
        return True
    except Exception as e:
        print(f"Error saving document {output_path}: {e}", file=sys.stderr)
        return False

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Inject a project into a Word Resume portfolio table.")
    parser.add_argument("--title", required=True, help="Title of the project")
    parser.add_argument("--tech", required=True, help="Tech stack used in the project")
    parser.add_argument("--desc", required=True, help="1-2 sentence description of the project")
    parser.add_argument("--column", default="auto", choices=["auto", "left", "right"], help="Target column ('left', 'right', or 'auto')")
    parser.add_argument("--position", default="last", choices=["first", "last"], help="Where to insert the project ('first' or 'last')")
    parser.add_argument("--input", default="../Aby_Kibru_Portfolio.docx", help="Path to base input resume template")
    parser.add_argument("--output", default="../Aby_Kibru_Portfolio_Updated.docx", help="Path to save the updated resume")
    
    args = parser.parse_args()
    
    success = inject_project(args.input, args.output, args.title, args.tech, args.desc, args.column, args.position)
    sys.exit(0 if success else 1)
